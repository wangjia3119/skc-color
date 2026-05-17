"""
批量色号转化工具
支持输入：Excel (.xlsx)、CSV (.csv)、纯文本 (.txt)、Word (.docx)、图片（OCR）
输出：Excel，包含原色号、SKC、TPG↔TCX 互转结果

用法：
  python batch_convert.py input.xlsx [色号列名] [--to TCX|TPG] [--out output.xlsx]
  python batch_convert.py input.csv  --to TCX
  python batch_convert.py input.txt  --to TPG
  python batch_convert.py input.docx --to TCX
  python batch_convert.py image.png  --to TPG    # 需要 pytesseract 或 easyocr
"""
import sys, csv, re, argparse, os
sys.path.insert(0, 'C:/Users/jiawa/skc_color')

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# 加载潘通对照表
_PANTONE_DB = {}
_PANTONE_BASE_DB = {}
_CSV_PATH = 'C:/Users/jiawa/skc_color/pantone_skc_v2.csv'
with open(_CSV_PATH, encoding='utf-8-sig') as f:
    for r in csv.DictReader(f):
        key = r['pantone'].strip().upper()
        _PANTONE_DB[key] = r
        base = key
        for sfx in ('TPG', 'TCX'):
            if base.endswith(sfx):
                base = base[:-len(sfx)]
                break
        if base not in _PANTONE_BASE_DB:
            _PANTONE_BASE_DB[base] = r

FAMILY_NAMES = {
    '0': '白/米白', '1': '黄', '2': '橙', '3': '红',
    '4': '粉',      '5': '紫', '6': '蓝', '7': '绿', '8': '棕', '9': '灰/黑'
}

# 潘通色号正则：XX-XXXXXX 或 XX-XXXXXXTPG/TCX
_PANTONE_RE = re.compile(r'\b(\d{2}-\d{4}(?:TPG|TCX)?)\b', re.IGNORECASE)


def _lookup(pantone_input: str) -> dict | None:
    """查找色号，支持 TPG/TCX/裸号"""
    key = pantone_input.strip().upper()
    if key in _PANTONE_DB:
        return _PANTONE_DB[key]
    if not key.endswith('TPG') and not key.endswith('TCX'):
        if key + 'TPG' in _PANTONE_DB:
            return _PANTONE_DB[key + 'TPG']
    base = key
    for sfx in ('TPG', 'TCX'):
        if base.endswith(sfx):
            base = base[:-len(sfx)]
            break
    return _PANTONE_BASE_DB.get(base)


def _get_aliases(result: dict) -> tuple[str, str]:
    """从查找结果推算 TPG / TCX 互转色号"""
    p = result['pantone'].upper()
    base = p
    for sfx in ('TPG', 'TCX'):
        if base.endswith(sfx):
            base = base[:-len(sfx)]
            break
    return base + 'TPG', base + 'TCX'


def _convert_one(pantone_input: str, to_series: str) -> dict:
    """转化单条色号"""
    inp = pantone_input.strip().upper()
    result = _lookup(inp)
    tpg, tcx = _get_aliases(result) if result else (None, None)
    target = (tpg if to_series == 'TPG' else tcx) if to_series and result else None
    return {
        'input':    inp,
        'found':    result is not None,
        'skc':      result['skc'] if result else '',
        'name':     result['name'] if result else '',
        'hex':      result['hex'] if result else '',
        'family':   FAMILY_NAMES.get(result['family'], '') if result else '',
        'tpg':      tpg or '',
        'tcx':      tcx or '',
        'target':   target or '',
        'review':   result['needs_review'] == 'True' if result else False,
    }


# ──────────────── 输入解析 ────────────────

def _read_excel(path: str, col: str) -> list[str]:
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    headers = [str(c.value).lower() if c.value else '' for c in ws[1]]
    try:
        idx = next(i for i, h in enumerate(headers) if col.lower() in h) + 1
    except StopIteration:
        # 没找到指定列，尝试全列扫描找潘通号
        print(f'未找到列 "{col}"，扫描全表提取潘通色号...')
        pantones = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            for cell in row:
                if cell:
                    for m in _PANTONE_RE.findall(str(cell)):
                        pantones.append(m)
        return pantones
    return [str(ws.cell(row=r, column=idx).value).strip()
            for r in range(2, ws.max_row + 1)
            if ws.cell(row=r, column=idx).value]


def _read_csv(path: str, col: str) -> list[str]:
    with open(path, encoding='utf-8-sig') as f:
        rows = list(csv.DictReader(f))
    if not rows:
        return []
    # 找列
    headers = list(rows[0].keys())
    key = next((h for h in headers if col.lower() in h.lower()), None)
    if not key:
        print(f'未找到列 "{col}"，扫描全行提取潘通色号...')
        pantones = []
        for row in rows:
            for v in row.values():
                for m in _PANTONE_RE.findall(str(v)):
                    pantones.append(m)
        return pantones
    return [row[key].strip() for row in rows if row.get(key, '').strip()]


def _read_txt(path: str) -> list[str]:
    with open(path, encoding='utf-8-sig') as f:
        text = f.read()
    return _PANTONE_RE.findall(text)


def _read_docx(path: str) -> list[str]:
    try:
        from docx import Document
    except ImportError:
        print('ERROR: 需要安装 python-docx：pip install python-docx')
        sys.exit(1)
    doc = Document(path)
    text = '\n'.join(p.text for p in doc.paragraphs)
    # 也扫描表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += '\n' + cell.text
    return _PANTONE_RE.findall(text)


def _read_image(path: str) -> list[str]:
    """OCR 提取图片中的潘通色号，优先用 easyocr，fallback pytesseract"""
    try:
        import easyocr
        reader = easyocr.Reader(['en'], verbose=False)
        results = reader.readtext(path, detail=0)
        text = ' '.join(results)
    except ImportError:
        try:
            import pytesseract
            from PIL import Image
            text = pytesseract.image_to_string(Image.open(path))
        except ImportError:
            print('ERROR: 需要安装 OCR 库：pip install easyocr  或  pip install pytesseract pillow')
            sys.exit(1)
    return _PANTONE_RE.findall(text)


def read_input(path: str, col: str = 'pantone') -> list[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext in ('.xlsx', '.xlsm'):
        return _read_excel(path, col)
    elif ext == '.csv':
        return _read_csv(path, col)
    elif ext == '.txt':
        return _read_txt(path)
    elif ext == '.docx':
        return _read_docx(path)
    elif ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp'):
        return _read_image(path)
    else:
        print(f'不支持的文件格式: {ext}')
        sys.exit(1)


# ──────────────── 输出 Excel ────────────────

def write_output(rows: list[dict], to_series: str, out_path: str):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '色号转化结果'

    hdr_fill = PatternFill('solid', fgColor='1A1A2E')
    hdr_font = Font(color='FFFFFF', bold=True)
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    to_label = to_series if to_series else 'TPG/TCX'
    headers = ['输入色号', 'SKC编码', '颜色名称', 'HEX', '色系', 'TPG色号', 'TCX色号', f'转化结果({to_label})', '颜色预览', '需审核', '状态']
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = Alignment(horizontal='center')

    for ri, row in enumerate(rows, 2):
        ws.cell(row=ri, column=1,  value=row['input'])
        ws.cell(row=ri, column=2,  value=row['skc'])
        ws.cell(row=ri, column=3,  value=row['name'])
        ws.cell(row=ri, column=4,  value=row['hex'])
        ws.cell(row=ri, column=5,  value=row['family'])
        ws.cell(row=ri, column=6,  value=row['tpg'])
        ws.cell(row=ri, column=7,  value=row['tcx'])
        ws.cell(row=ri, column=8,  value=row['target'])
        # 颜色预览
        preview = ws.cell(row=ri, column=9, value='  ')
        if row['hex']:
            preview.fill = PatternFill('solid', fgColor=row['hex'].lstrip('#').upper())
        ws.cell(row=ri, column=10, value='是' if row['review'] else '否')
        status = '已找到' if row['found'] else '未找到'
        status_cell = ws.cell(row=ri, column=11, value=status)

        # SKC 格式高亮
        skc_cell = ws.cell(row=ri, column=2)
        if not row['found']:
            status_cell.font = Font(color='B71C1C')
            skc_cell.font = Font(color='999999')
        elif row['review']:
            skc_cell.fill = PatternFill('solid', fgColor='FFEB3B')
            skc_cell.font = Font(bold=True, color='B71C1C')
        else:
            skc_cell.font = Font(bold=True, color='1B5E20')

        for ci in range(1, len(headers) + 1):
            ws.cell(row=ri, column=ci).border = border

    col_widths = [16, 10, 22, 10, 10, 16, 16, 16, 10, 8, 8]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    # 统计行
    total = len(rows)
    found = sum(1 for r in rows if r['found'])
    review = sum(1 for r in rows if r['review'])
    ws.cell(row=total + 3, column=1, value=f'共 {total} 条，找到 {found} 条，未找到 {total-found} 条，需审核 {review} 条').font = Font(italic=True, color='666666')

    wb.save(out_path)
    print(f'已写入: {out_path}')
    print(f'共 {total} 条 | 找到 {found} | 未找到 {total-found} | 需审核 {review}')


# ──────────────── 主程序 ────────────────

def main():
    parser = argparse.ArgumentParser(description='潘通色号批量转化 (TPG ↔ TCX)')
    parser.add_argument('input', help='输入文件路径（xlsx/csv/txt/docx/图片）')
    parser.add_argument('col', nargs='?', default='pantone', help='色号列名（Excel/CSV，默认 pantone）')
    parser.add_argument('--to', default='', choices=['TPG', 'TCX', ''], help='目标系列')
    parser.add_argument('--out', default='', help='输出文件路径')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f'ERROR: 文件不存在: {args.input}')
        sys.exit(1)

    pantones = read_input(args.input, args.col)
    # 去重保序
    seen = set()
    pantones = [p for p in pantones if p and not (p.upper() in seen or seen.add(p.upper()))]
    print(f'提取到 {len(pantones)} 个色号')

    if not pantones:
        print('未提取到任何潘通色号，退出。')
        sys.exit(0)

    to_series = args.to.upper()
    results = [_convert_one(p, to_series) for p in pantones]

    out_path = args.out or args.input.rsplit('.', 1)[0] + '_converted.xlsx'
    write_output(results, to_series, out_path)


if __name__ == '__main__':
    main()
